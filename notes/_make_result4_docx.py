"""Generate a standalone Result 4 draft as a .docx in notes/.

Contents:
- Result 4 prose (6 paragraphs)
- Methods subsections relevant to Result 4 (Specialist models, Encoder PCA)

Style rules:
- No em dashes anywhere (use commas, periods, parentheses).
- No colons or semicolons in body prose. List-introductions become full sentences.
- Helvetica 11 pt, 1-inch margins, first-line indent on body paragraphs.
- Subsection headings in bold.
- Don't cite Whole Root as a per-class metric; numbers use 6-class mIoU
  (excludes Whole Root).
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/yifeigu/Documents/Siobhan_Lab/plants/notes/result4_revised_2026_05_10.docx"

SECTIONS = [
    (
        "Result 4. A unified model for monocots and dicots generalizes better than "
        "the specialist models trained on monocots or dicots individually",
        [
            "Cereal cross-section images (rice, sorghum, pearl millet) and tomato "
            "cross-section images differ substantially in cell-wall architecture, "
            "root diameter, and the presence of aerenchyma. A natural concern when "
            "training a single model jointly on both clades is that the model "
            "might dilute its representation of either clade in service of the "
            "other, and that a clade-specific specialist would perform better on "
            "its own clade than a generalist. To test this directly, we trained "
            "two additional models. The Monocot Specialist was trained only on "
            "the 856 monocot training samples (rice, sorghum, and pearl millet). "
            "The Dicot Specialist was trained only on the 437 tomato training "
            "samples. The Unified Model and the two specialists shared an "
            "identical encoder, decoder, loss combination, augmentation pipeline, "
            "optimizer, and schedule, with the only difference being the species "
            "composition of the training set.",

            "We evaluated all three models on three test cohorts. The cohorts "
            "comprised the 148 monocot in-distribution test samples, the 37 dicot "
            "in-distribution test samples, and the 35 rice samples held out on "
            "the Zeiss confocal microscope. Representative prediction examples "
            "from each model on a monocot test sample (Millet, AW23 genotype), a "
            "dicot test sample (Tomato, slmyb92 mutant), and a Zeiss zero-shot "
            "sample (Rice, Kittake) are shown in Fig. 3a alongside the input "
            "image and the expert ground truth. The quantitative comparison "
            "across all 220 evaluation samples is summarized in Fig. 3b. On its "
            "own training clade, each specialist matched the Unified Model. On "
            "the 148 monocot test samples, the Monocot Specialist reached "
            "6-class mIoU 0.856 and the Unified Model reached 0.857, a "
            "difference of 0.001. On the 37 dicot test samples, the Dicot "
            "Specialist reached 6-class mIoU 0.932 and the Unified Model reached "
            "0.936, a difference of 0.004. On the 35 Zeiss zero-shot rice "
            "samples, both monocot-trained models reached 6-class mIoU 0.873 "
            "(Unified) and 0.872 (Monocot Specialist), again a difference of "
            "0.001. Joint training across clades therefore preserves full "
            "specialist-level accuracy on each clade rather than diluting it.",

            "The picture changed dramatically when each specialist was evaluated "
            "on the clade it had never seen during training (Fig. 3b, off-domain "
            "bars). The Monocot Specialist, applied to the 37 dicot test "
            "samples, dropped from 6-class mIoU 0.856 on its own monocot test "
            "set to 0.603 on tomato, a fall of 0.253. The Dicot Specialist, "
            "applied to the 148 monocot test samples, dropped from 0.932 on its "
            "own dicot test set to 0.589 on cereals, a fall of 0.343. On the "
            "Zeiss zero-shot split, which consists entirely of rice samples and "
            "is therefore off-domain for the Dicot Specialist, the Dicot "
            "Specialist collapsed further to 6-class mIoU 0.401, an absolute "
            "drop of 0.531 from its dicot in-distribution accuracy. By "
            "contrast, the Unified Model maintained 6-class mIoU between 0.857 "
            "and 0.936 across every cohort. The visual prediction panels in "
            "Fig. 3a illustrate this failure mode directly. The off-clade "
            "specialist predictions either miss entire tissue domains or "
            "misclassify large regions as the wrong tissue, while the Unified "
            "Model's prediction is essentially indistinguishable from the "
            "expert ground truth on both clades. Specialist training therefore "
            "produces a brittle model that fails catastrophically when applied "
            "beyond its training clade, whereas joint training produces a "
            "generalist model with no measurable performance cost.",

            "We also asked whether fine-tuning on our annotated dataset "
            "reshapes the encoder's internal representation in a way that "
            "reflects root anatomy, rather than only modifying the downstream "
            "segmentation head. To answer this, we visualized the spatial "
            "structure of the encoder's learned representations directly on "
            "the input image (Fig. 3c). For each sample, we passed the image "
            "through the encoder and extracted the grid of patch-token "
            "feature vectors output by the final transformer block, then "
            "projected these high-dimensional vectors into a three-dimensional "
            "colour space using principal component analysis (PCA). PCA is a "
            "standard linear dimensionality-reduction method that identifies "
            "the directions of greatest variance in a high-dimensional "
            "feature set and projects the data onto those directions. When "
            "applied to patch-token features, PCA recovers the dominant "
            "patterns of variation across patches within an image. We "
            "rendered the first three principal components as the red, green, "
            "and blue channels of an output image at the patch grid "
            "resolution, with PCA computed separately for the pretrained "
            "encoder and for the fine-tuned Unified Model encoder (full "
            "procedure in Methods).",

            "Fig. 3c shows the encoder PCA visualization for three "
            "representative samples, namely a monocot test image (Rice, "
            "wox10-15 mutant), a dicot test image (Tomato, slmyb92 mutant), "
            "and a Zeiss zero-shot monocot image (Rice, Kittake), each "
            "rendered under two encoder states. The first column shows the "
            "pretrained DINOv3 encoder before any fine-tuning on our dataset, "
            "and the second column shows the fine-tuned Unified Model encoder. "
            "In the pretrained encoder, the PCA colours form diffuse and "
            "largely amorphous patterns within the root that do not align with "
            "the underlying anatomical tissue domains. In the Unified Model "
            "encoder, by contrast, the PCA colours form sharp, spatially "
            "coherent rings and regions that align directly with the "
            "epidermis, cortex, endodermis, exodermis, and vascular cylinder "
            "visible in the input image. Concentric ring structure emerges in "
            "the encoder feature space that mirrors the radial anatomical "
            "structure of the root, and aerenchyma regions are picked out as "
            "distinct colour patches embedded within the cortex. This pattern "
            "holds across all three rows, including the Zeiss zero-shot row in "
            "which the encoder had never seen confocal imaging optics during "
            "training.",

            "This comparison confirms that fine-tuning on our annotated root "
            "dataset successfully reshapes the encoder's representation. "
            "Before fine-tuning, the pretrained DINOv3 encoder produces a "
            "feature space whose dominant axes of variation do not correspond "
            "to root anatomical structure. After fine-tuning, the encoder "
            "learns a more structured feature space that is correlated with "
            "the radial anatomical organization of the root, with each tissue "
            "domain occupying a distinct region of the feature space. The "
            "encoder is therefore not merely supporting a downstream "
            "segmentation head, it is learning representations that are "
            "themselves anatomically meaningful.",
        ],
    ),

    (
        "Methods subsections relevant to Result 4",
        [
            {"sub": "Specialist models (Monocot Specialist and Dicot Specialist)"},
            "Two clade-specialist models were trained for the unified-versus-"
            "specialist comparison in Result 4. The Monocot Specialist was "
            "trained on the 856 training samples drawn from rice, sorghum, and "
            "pearl millet. The Dicot Specialist was trained on the 437 training "
            "samples drawn from tomato (including the cultivated Solanum "
            "lycopersicum and 13 wild Solanum species). The validation set was "
            "restricted to the matching clade for each specialist. Both "
            "specialists used the same DINOv3-S/16 encoder, Meta-style DPT "
            "decoder, four-component composite loss, augmentation pipeline, "
            "optimizer, learning-rate schedule, batch size, and epoch budget "
            "as the Unified Model, with the only difference being the "
            "training-set species composition. Specialists were evaluated on "
            "three cohorts. Each was evaluated on its own clade's "
            "in-distribution test split using the standard "
            "metrics_bio7.csv pipeline. Each was also evaluated on the "
            "off-clade test split (the Monocot Specialist on the 37 tomato "
            "test samples, the Dicot Specialist on the 148 cereal test "
            "samples) using a separate off-domain inference pass written to a "
            "metrics_bio7_offdomain.csv file. Finally, both specialists were "
            "applied zero-shot to the 35 Zeiss rice samples for direct "
            "comparison with the Unified Model on the zero-shot split.",

            {"sub": "Encoder feature visualization by principal component analysis"},
            "To visualize the spatial organization of the encoder's learned "
            "representations on a per-sample basis, we used a two-pass "
            "principal component analysis (PCA) procedure adapted from the "
            "DINOv2 and DINOv3 feature-visualization conventions (Oquab et al. "
            "2023, Sim\u00e9oni et al. 2025). For each sample, we resampled "
            "the input image to 1,536 \u00d7 1,536 pixels (slightly above the "
            "standard 1,024 \u00d7 1,024 training resolution to give the "
            "visualization a finer patch grid), passed it through the chosen "
            "encoder (either the frozen pretrained DINOv3-S/16 or the "
            "fine-tuned Unified Model encoder), and extracted the patch-token "
            "output of the final transformer block. With a patch size of 16, "
            "this produced a 96 \u00d7 96 grid of 384-dimensional patch-token "
            "feature vectors per sample.",

            "The first PCA pass fit a one-component PCA on the patch tokens of "
            "a single image. The first principal component of self-supervised "
            "DINO-family features is known to separate foreground (object) "
            "from background (surroundings) without supervision. We "
            "thresholded the PC1 scores at their median to obtain a binary "
            "foreground mask, choosing the foreground side as the one with "
            "fewer image-corner patches under the assumption that the four "
            "corners of a cross-section image are background.",

            "The second PCA pass fit a three-component PCA on the union of "
            "foreground tokens pooled across all samples for the encoder "
            "being visualized, with each sample's foreground tokens "
            "mean-centered before pooling. The pretrained DINOv3 encoder and "
            "the fine-tuned Unified Model encoder used separate PCA bases, "
            "each fit only on its own foreground tokens, because the two "
            "encoders live in different feature spaces and a shared PCA "
            "basis would not be meaningful between them. For each (encoder, "
            "sample) pair, the three component scores at each foreground "
            "patch were clipped at their per-component 2nd and 98th "
            "percentile values (pooled across the fit), then linearly "
            "rescaled to the unit interval and assigned to the red (PC1), "
            "green (PC2), and blue (PC3) channels of an output image at the "
            "patch grid resolution. Background patches were rendered black. "
            "The result is a per-sample colour visualization in which "
            "patches with similar encoder representations appear in similar "
            "colours, so the spatial layout of those colours directly "
            "reflects the structure encoded in the encoder's feature space.",
        ],
    ),
]

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Helvetica"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Result 4 + related Methods (revised draft, 2026-05-10)")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

audit_text_chunks = []

for heading, items in SECTIONS:
    h = doc.add_paragraph()
    h_run = h.add_run(heading)
    h_run.bold = True
    h_run.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)

    for item in items:
        if isinstance(item, dict) and "sub" in item:
            sh = doc.add_paragraph()
            sh_run = sh.add_run(item["sub"])
            sh_run.bold = True
            sh_run.font.size = Pt(11)
            sh.paragraph_format.space_before = Pt(8)
            sh.paragraph_format.space_after = Pt(2)
            audit_text_chunks.append(item["sub"])
        else:
            body = doc.add_paragraph(item)
            body.paragraph_format.first_line_indent = Inches(0.25)
            body.paragraph_format.space_after = Pt(8)
            audit_text_chunks.append(item)

audit_text = "\n".join(audit_text_chunks)
banned = {"\u2014": "em dash", ";": "semicolon", ":": "colon"}
for ch, name in banned.items():
    if ch in audit_text:
        idx = audit_text.index(ch)
        ctx = audit_text[max(0, idx - 50) : idx + 50]
        raise AssertionError(f"{name} found in body prose near: ...{ctx}...")

doc.save(OUT)
n_words = sum(
    len(item.split()) for _, items in SECTIONS for item in items
    if isinstance(item, str)
)
print(f"Wrote {OUT}")
print(f"Word count: {n_words}")
