"""Generate the revised introduction as a .docx in notes/."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/yifeigu/Documents/Siobhan_Lab/plants/notes/intro_revised_2026_05_04.docx"

# (label, kind, text). kind in {"verbatim", "new", "edited"}.
PARAS = [
    ("Paragraph 1", "verbatim",
     "Roots are considered the hidden half of plants and their presence below-ground historically "
     "precluded facile understanding of their form and function. While high-throughput acquisition "
     "of root system architecture images and their subsequent automated quantification have vastly "
     "improved, anatomical (cellular) quantification remains more difficult. Plant root cell types "
     "are radially symmetrical and arranged in concentric cylinders surrounding the vasculature. "
     "The epidermis is the outer cell layer, followed by the cortex and endodermis. The root stele "
     "is composed of the pericycle, as well as the xylem, phloem and procambial tissue. Despite "
     "conservation of this patterning across plant species, the number of layers and their "
     "morphology are phylogenetically variable. Arabidopsis thaliana, frequently used as a model "
     "species, contains a single cortex layer, unlike most plant species that contain multiple "
     "cortex layers, many of which are specialized."),

    ("Paragraph 2", "verbatim",
     "Functional specialization can be conferred by deposition of a complex cell wall composed of "
     "lignin, suberin, or other biopolymers. Deposition of these walls controls water and mineral "
     "ion uptake, blocks pathogens, or its removal can facilitate entry of symbionts (Naseer et al. "
     "2012; Baxter et al. 2009; Cantó-Pastor et al. 2024; Manzano et al. 2025; Barberon et al. "
     "2016; Reynoso et al. 2022; Zhu et al. 2025). The outer cortical cell type, the exodermis, "
     "can be lignified or suberized, and inner cortical layers can contain lignin depositions of "
     "various types (Schneider et al. 2021)(sclerenchyma + phi thickening references). The root "
     "endodermis contains a conserved lignified Casparian Strip, and can be suberized (Naseer et "
     "al. 2012; Barberon et al. 2016). Monocot root species can contain aerenchyma, areas of cell "
     "wall loosening and programmed cell death, producing gaps in the cellular structure that can "
     "reduce metabolic cost or maintain oxygen in stressful environments (Abiko et al. 2012; Klein "
     "et al. 2020). All of these cell type features are dynamically modified in response to "
     "abiotic and biotic stressors, and their increased presence is often associated with stress "
     "tolerance (Gouran and Brady 2024). A central question remains, what is the diversity of "
     "these cell types in plants, their differentiation features, and how do they enable adaptive "
     "plant responses to the environment?"),

    ("Paragraph 3", "verbatim",
     "Anatomical resolution of these cell types, cell wall barriers and aerenchyma require thin "
     "sections and cellular visualization. Laser ablation tomography is an effective method to "
     "determine cortical cell wall features, and its application revealed the genetic architecture "
     "of these traits and their functional and environmental significance in maize, wheat and "
     "their wild relatives, but is not an easily accessible technology (Schneider et al. 2023, "
     "2021; Klein et al. 2020). Sections can be simply obtained using vibratomes and fresh or "
     "agar-embedded specimens (Sexauer et al. 2021; Manzano et al. 2025; Cantó-Pastor et al. 2024; "
     "Kawa et al. 2024). Establishment of the ClearSee protocol enabled deeper penetration of "
     "root tissues, and coupled with fluorescent histochemical dyes has already revealed a "
     "complex and rich spectrum of these structures in a small number of species (Sexauer et al. "
     "2021; Ursache et al. 2018; Kurihara et al. 2015). These methods now make it straightforward "
     "to generate large collections of high-resolution root images across species, genotypes, and "
     "experimental conditions. However, the quantitative analysis of these images remains a major "
     "bottleneck, limiting the scale and scope of biological studies that can be performed."),

    ("Paragraph 4", "verbatim",
     "Traditionally, quantification of root barriers relies on categorical predictors (Barberon et "
     "al. 2016; Naseer et al. 2012) or manual annotation (Kawa et al. 2024), a process that is "
     "time-intensive and difficult to scale, particularly when employed for quantitative genetic "
     "studies. This challenge is particularly pronounced for anatomically complex structures such "
     "as cortical aerenchyma, where boundaries can be irregular, diffuse, or discontinuous, making "
     "annotation subjective and prone to inter- and intra-annotator variability. As a result, "
     "datasets that require months of effort to analyze are often restricted to small sample "
     "sizes, and rich image collections are underutilized."),

    ("Paragraph 5", "new",
     "From a computer vision perspective, root anatomy is not a collection of independent objects "
     "but a nested, semantically organized tissue map. The identity of a region depends as much "
     "on its global position within the root as on its local cellular appearance. The exodermis, "
     "for example, is operationally defined as the second cell layer from the outer surface. In "
     "some sorghum samples, it is heavily lignified or suberized and produces a strong, "
     "distinguishing signal in the FITC and TRITC channels. In other samples of the same species, "
     "it is unlignified, unsuberized, and visually indistinguishable from the surrounding "
     "cortical cells. A model that classifies regions from local texture alone cannot recover the "
     "correct label in the second case. Robust segmentation therefore requires a model that "
     "learns the semantic topology of the root, including the ordered radial arrangement of "
     "epidermis, exodermis, cortex, endodermis, vasculature, and aerenchyma, and that integrates "
     "long-range context to resolve regions whose local appearance is ambiguous."),

    ("Paragraph 6", "new",
     "Existing biological segmentation tools were not designed for this task. Cellpose, PlantSeg, "
     "MorphoGraphX (Barbier de Reuille et al. 2015), and the Segment Anything Model are "
     "class-agnostic instance segmenters. They localize individual cells and trace cell "
     "boundaries, but they carry no representation of which tissue a cell belongs to, and their "
     "decisions are made from a local receptive field that does not capture radial position. "
     "Class assignment must then be performed manually or through a separate classifier. A "
     "second limitation is geometric. Watershed-style pipelines, which are used by Cellpose and "
     "many cell-segmentation packages, place the boundary between two touching cells at the "
     "midline between their centers. This convention splits the shared cell wall in half. When "
     "the analyte of interest is deposited in the wall itself, as is the case for suberin and "
     "lignin, the watershed approach discards roughly half of the signal. Plant root anatomy "
     "instead requires the opposite behavior: a semantic segmentation in which a tissue region, "
     "for example the entire exodermis ring, includes the full thickness of its bounding cell "
     "walls."),

    ("Paragraph 7", "new",
     "Two recent advances in computer vision make this problem newly tractable. First, "
     "self-supervised vision foundation models, most recently DINOv3 trained on the LVD-1.69B "
     "image collection, produce dense patch features that encode semantic and topological "
     "structure. These features transfer across image domains, including microscopy, with minimal "
     "fine-tuning. Second, dense prediction transformer (DPT) decoders fuse features from "
     "multiple encoder layers into pixel-level outputs while preserving long-range context, which "
     "makes them well suited to nested anatomical maps. When fine-tuned on a sufficiently diverse "
     "annotated dataset, a foundation-model encoder paired with a DPT decoder can deliver the "
     "contextual reasoning that root anatomy requires."),

    ("Paragraph 8", "new",
     "Here we test this proposition. We introduce a multi-species, multi-platform dataset of "
     "approximately 1,600 expert-annotated root cross-section images. The dataset covers both "
     "monocots (rice, sorghum, and millet) and a dicot (tomato), three imaging platforms "
     "(Olympus widefield, Cytation C10, and Zeiss confocal), multiple genotypes, and multiple "
     "growth conditions. We focus on the key cell wall barriers of the endodermis and exodermis, "
     "together with cortical aerenchyma, which regulate radial transport, nutrient and water "
     "flux, gas exchange, and plant responses to environmental stress and biotic interactions "
     "(Geldner 2013; Lynch 2015; Barberon 2017; Yamauchi et al. 2018). On this dataset we train "
     "a unified semantic segmentation model based on a fine-tuned DINOv3-S/16 encoder paired "
     "with a Meta-style DPT decoder. The model simultaneously segments seven tissue classes, "
     "including the lignified and suberized barriers of the endodermis and exodermis as well as "
     "cortical aerenchyma. The unified model achieves per-class IoU within the inter-annotator "
     "ceiling estimated from a second expert (Fig. 2), outperforms species- and clade-specialist "
     "models trained on the same data (Fig. 3), and transfers zero-shot to an entirely held-out "
     "imaging platform. Its predictions are accurate enough to substitute for manual annotation "
     "in three quantitative biological readouts: endodermal suberin intensity, exodermal lignin "
     "intensity, and aerenchyma area fraction. The resulting measurements match expert ground "
     "truth across the test set (Fig. 4). Finally, we deploy the model at scale to "
     "[Fig 5 biological finding, to be filled in], showing that automated anatomical "
     "quantification can resolve genotype by environment effects on root barrier architecture "
     "at a throughput previously inaccessible to the field. Together, the dataset, the model, "
     "and the downstream pipeline provide a scalable framework for dissecting how root barrier "
     "architecture contributes to plant responses to environmental and biotic stress."),
]

KIND_LABEL = {
    "verbatim": "verbatim from biologist co-author",
    "new":      "new (Yifei, AI/CV perspective)",
    "edited":   "edited (kept biologist opening, extended with AI/CV content)",
}
KIND_COLOR = {
    "verbatim": RGBColor(0x55, 0x55, 0x55),
    "new":      RGBColor(0x1f, 0x66, 0x9e),
    "edited":   RGBColor(0x9e, 0x66, 0x1f),
}

doc = Document()

# Page setup: 1 inch margins, body in 11 pt Helvetica.
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Helvetica"
style.font.size = Pt(11)

# Title.
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Introduction (revised draft, 2026-05-04)")
run.bold = True
run.font.size = Pt(14)

# Provenance key.
key = doc.add_paragraph()
key_run = key.add_run(
    "Source key: italic grey heading above each paragraph indicates whether the paragraph is "
    "verbatim from the biologist co-author, edited (biologist opening kept, AI/CV content added), "
    "or new (Yifei). The headings are for review only and would be removed before submission."
)
key_run.italic = True
key_run.font.size = Pt(9)
key_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()  # spacer

for label, kind, text in PARAS:
    head = doc.add_paragraph()
    head_run = head.add_run(f"{label}  —  {KIND_LABEL[kind]}".replace("—", "|"))
    head_run.italic = True
    head_run.font.size = Pt(9)
    head_run.font.color.rgb = KIND_COLOR[kind]

    body = doc.add_paragraph(text)
    body.paragraph_format.first_line_indent = Inches(0.25)
    body.paragraph_format.space_after = Pt(8)

# Sanity check: no em dashes anywhere in body text.
all_text = "\n".join(t for _, _, t in PARAS)
assert "—" not in all_text, "em dash slipped into the body"

doc.save(OUT)
print(f"Wrote {OUT}")
print(f"Body word count: {sum(len(t.split()) for _,_,t in PARAS)}")
