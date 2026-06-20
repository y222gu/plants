"""Generate the revised Results sections (1 and 2) plus a Methods draft as a
.docx in notes/.

Style rules in this draft:
- No em dashes anywhere (use commas, periods, parentheses).
- No colons or semicolons in body prose. List-introductions become full sentences.
- Helvetica 11 pt, 1-inch margins, first-line indent on body paragraphs.
- Methods subsection headings in bold, 11 pt.
- Biology subsections that need co-author input are written as italic
  placeholder paragraphs starting with "[TBD by co-author]".

Update 2026-05-10:
- Result 2 restructured with corrected Fig 2 panel letters (a, b = dicot/monocot
  galleries; c = Zeiss zero-shot gallery; d = per-class boxplot; e = per-microscope
  boxplot; f = architecture comparison scatter).
- New segmentation-example paragraph (Fig 2a, b).
- New failure-mode paragraph (Rice Wox10-50 outlier cluster + small aerenchyma).
- Fusion-mechanism paragraph removed from Result 2 (full mechanism stays in Methods).
- Methods Evaluation subsection now opens with explicit IoU and Dice definitions.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/yifeigu/Documents/Siobhan_Lab/plants/notes/results_revised_2026_05_10.docx"

PLACEHOLDER = RGBColor(0x99, 0x55, 0x33)

# ── Content tree ────────────────────────────────────────────────────────────
SECTIONS = [
    # ── Result 1 ───────────────────────────────────────────────────────────
    (
        "A multi-species fluorescence microscopy dataset of cereal and dicot root "
        "cross-sections annotated for six anatomical tissue domains",
        [
            "Supervised semantic segmentation of plant root anatomy is bottlenecked by the "
            "scale and breadth of expert-annotated training data. Existing fluorescence "
            "microscopy datasets are typically restricted to a single species, a single "
            "imaging platform, or a single laboratory, which limits the generalization of "
            "models trained on them and forces each new study to retrain from scratch. To "
            "overcome this bottleneck, we assembled a benchmark dataset of 1,695 root "
            "cross-section images spanning four species and three fluorescence imaging "
            "platforms, annotated by plant biologists for six anatomical tissue domains "
            "(Fig. 1a). The dataset was designed not simply to be large, but to vary "
            "deliberately along the three axes that determine whether a trained model can "
            "be deployed beyond the laboratory in which it was trained. These axes are "
            "species, genotype, and imaging hardware.",

            "The dataset is dominated by monocot cereals, with 1,150 samples drawn from "
            "rice, sorghum, and pearl millet, and is balanced by 545 dicot tomato "
            "cross-sections that together provide both monocot and dicot anatomy and "
            "force the model to learn tissue features that generalize across the major "
            "anatomical division among flowering plants. Within each species, samples "
            "span multiple genotypes contributed by laboratories at several institutions "
            "and geographic regions (Fig. 1b), capturing biological variation in "
            "cell-wall composition, root diameter, and aerenchyma development. This "
            "breadth is essential. Models trained on a single genotype tend to overfit "
            "to that genotype's idiosyncratic anatomy, and a model that is to be useful "
            "across the cereal research community must recognize the same tissue domains "
            "across markedly different root morphologies.",

            "Each cross-section was stained with three fluorescent dyes that label "
            "distinct cell-wall constituents. DAPI provided a general cell-wall "
            "counterstain, FITC labelled lignin, and TRITC labelled suberin. To capture "
            "the optical and noise variation introduced by different imaging set-ups, "
            "samples were imaged on three platforms. The Olympus IX83 widefield "
            "microscope contributed 1,501 samples, the Cytation C10 plate-reader imager "
            "contributed 159 samples, and the Zeiss LSM 970 confocal microscope "
            "contributed 35 rice samples. All images were resampled to 1,024 \u00d7 "
            "1,024 pixels for model input but retained the platform-specific spatial "
            "resolution, signal-to-noise characteristics, and background features of "
            "the original acquisition (Fig. 1c). In practice, different research groups "
            "image their root sections on whatever microscope is locally available. "
            "Including three optical regimes within the dataset directly tests whether "
            "a single trained model can absorb that variation rather than requiring "
            "each user to retrain on their own platform.",

            "Each sample was annotated by trained plant biologists with hand-drawn "
            "polygons in a custom interactive editor (full annotation protocol in the "
            "Supplementary Information). Polygons conformed to a fixed cell-wall-boundary "
            "convention, with ring tissues delineated along the outer edge of their "
            "respective cell walls, ensuring consistency across annotators, samples, and "
            "imaging platforms. Across the 1,695 samples, annotators drew a total of "
            "31,595 polygons. These comprised 1,695 whole-root contours, 3,390 "
            "endodermal ring contours (one outer and one inner per sample), 3,390 "
            "exodermal ring contours, and 23,119 aerenchyma polygons. Aerenchyma counts "
            "ranged from zero per section in tomato (which lacks aerenchyma) to dozens "
            "of individual lacunae in mature cereal sections. These polygons jointly "
            "define the six anatomical tissue domains used as model targets, namely "
            "epidermis, cortex, endodermis, exodermis, vascular cylinder, and "
            "aerenchyma. To our knowledge, this represents the most extensive annotated "
            "fluorescence microscopy dataset of plant root cross-sections currently "
            "available, with the per-tissue resolution required to support quantitative "
            "trait analyses that go beyond simple region-of-interest detection.",

            "We partitioned the dataset into four splits, each designed to probe a "
            "distinct facet of model robustness. The training set (1,293 samples) was "
            "used for parameter optimization, and the validation set (182 samples) was "
            "used for model selection and early stopping. The remaining samples were "
            "divided into two complementary test sets that target different "
            "generalization regimes. The first, an in-distribution test set (185 "
            "samples), contained cross-sections drawn from the same four species and "
            "the same two widefield platforms as the training data (Olympus IX83 and "
            "Cytation C10), but from biological experiments that were entirely held out "
            "from training. This split tests whether the model has learned the "
            "underlying anatomy of root tissues rather than memorizing experiment-"
            "specific cues such as a particular growth condition, lighting, or "
            "sectioning artifact, and it quantifies the performance a researcher should "
            "expect when applying the model to new samples acquired under their "
            "existing imaging conditions. The second, a zero-shot test set (35 rice "
            "samples), consisted of cross-sections imaged on an entirely different "
            "platform that the model never saw during training, the Zeiss LSM 970 "
            "confocal microscope. This split simulates the most demanding deployment "
            "scenario, in which a new laboratory adopts the model for use on its own "
            "microscopy hardware, and it quantifies the model's robustness to optical "
            "and acquisition variation that goes beyond the diversity sampled in the "
            "training set. Together, these two test sets bracket the failure modes that "
            "matter for real-world deployment. The in-distribution split reveals "
            "whether the model has overfit to specific experiments, and the zero-shot "
            "split reveals whether it has overfit to the specific imaging platforms it "
            "has seen. To prevent experiment-level information from leaking between "
            "training and testing, we required all samples from a single experiment to "
            "remain within a single split.",
        ],
    ),

    # ── Result 2 ───────────────────────────────────────────────────────────
    (
        "RootSeg outperforms human annotators and other root anatomical semantic "
        "segmentation models",
        [
            # Para 1: model overview (brief fusion mention only; mechanism in Methods)
            "Building on this large, high-quality dataset, we developed RootSeg, a "
            "deep-learning model that takes a three-channel fluorescence image as "
            "input and returns a per-pixel cell type identity map. The model has two "
            "main components (Fig. 1d). The encoder is a vision transformer "
            "initialized from DINOv3, a self-supervised foundation model pretrained "
            "by Meta on the LVD-1.69B natural-image corpus [ref]. It compresses the "
            "image into a learned numerical representation that captures both local "
            "cellular detail and the global context of the root. The decoder is a "
            "dense prediction transformer (DPT), which expands that representation "
            "back into a tissue map at full image resolution by fusing encoder "
            "features at multiple resolutions to combine local cellular detail with "
            "global anatomical context (full architecture in Methods).",

            # Para 2: topic sentence + in-distribution numbers (corrected Fig refs)
            "RootSeg achieved consistently high segmentation accuracy across all "
            "seven anatomical tissue classes. On the in-distribution test set of 185 "
            "samples spanning ten plant species across both monocot and dicot clades "
            "and the two widefield platforms used during training, RootSeg achieved "
            "a mean intersection-over-union (mIoU) of 0.889 and a mean Dice "
            "coefficient of 0.929 across the seven tissue classes (Fig. 2f). "
            "Performance was high across most anatomical compartments. Cortex and "
            "vascular cylinder each exceeded IoU 0.97, the endodermal and exodermal "
            "rings reached IoU 0.91 and 0.88, respectively, and the epidermis "
            "reached IoU 0.82 (Fig. 2d). The hardest class was aerenchyma at IoU "
            "0.60. This reflected both the small per-pixel footprint of individual "
            "lacunae and the high inter-species variation in aerenchyma morphology, "
            "with sparse, micron-scale lacunae in pearl millet contributing most of "
            "the residual error.",

            # Para 3 NEW: segmentation example panels
            "Representative segmentation examples (Fig. 2a, b) show the model's "
            "output on dicot and monocot test samples respectively, alongside the "
            "corresponding expert ground-truth annotations. Each example was "
            "selected to span the model's typical performance range. The dicot "
            "panel (Fig. 2a) includes one Solanum lycopersicum M82 wild-type sample "
            "on the C10 platform (mIoU 0.946), one suberin-mutant sample on Olympus "
            "(mIoU 0.913), and one wild Solanum sample contributed by an external "
            "collaborator (mIoU 0.941), illustrating that performance is preserved "
            "across cultivar background and across the two widefield platforms. The "
            "monocot panel (Fig. 2b) includes one sample per cereal-platform "
            "combination, with mIoU values ranging from 0.882 (pearl millet on "
            "Olympus) to 0.927 (sorghum on Olympus), and shows that the model "
            "resolves the thin endodermal and exodermal rings as well as the "
            "irregular shapes of individual aerenchyma lacunae across all three "
            "cereal species. Across both galleries, the predicted tissue maps are "
            "visually nearly indistinguishable from expert annotation.",

            # Para 4: per-species + per-microscope (Fig 2d species, Fig 2e platform)
            "Performance was consistent across species and imaging platforms within "
            "the in-distribution test set, indicating that a single unified model "
            "generalizes across the dataset's biological and optical heterogeneity. "
            "Per-species mean IoU was 0.945 for tomato, 0.898 for sorghum, 0.863 "
            "for rice, and 0.856 for pearl millet (Fig. 2d). Per-platform mean IoU "
            "was 0.918 on the Cytation C10 plate-reader imager and 0.885 on the "
            "Olympus IX83 widefield microscope (Fig. 2e). The slightly lower "
            "performance on monocots reflects the difficulty of the aerenchyma "
            "class, which is unique to monocots and contributes most of the "
            "residual classification error in those species.",

            # Para 5 NEW: failure modes
            "To understand the residual error, we examined the lowest-scoring "
            "samples in the in-distribution test set. Mean IoU per sample ranged "
            "from 0.608 to 0.949, with the bottom decile concentrated in two "
            "distinct failure modes. The first and dominant failure mode affected "
            "the thin outer rings, namely epidermis and exodermis. Across all 185 "
            "test samples, median epidermis IoU was 0.858 but the 10th percentile "
            "dropped to 0.713, and median exodermis IoU was 0.936 but the 10th "
            "percentile dropped to 0.820. Strikingly, nine of the ten lowest-mIoU "
            "samples came from a single rice experiment (Exp10, Wox10-50 mutant "
            "background imaged on Olympus), where epidermis IoU collapsed to "
            "between 0.12 and 0.44 and exodermis IoU to between 0.04 and 0.25, "
            "while cortex, endodermis, and vascular cylinder all remained above "
            "IoU 0.77. Visual inspection (Supp "
            "Fig. X) shows that in these samples the outermost cell layers were "
            "poorly stained or partially detached during sectioning, leaving the "
            "model with little local cellular signal to demarcate the epidermis "
            "from the exodermis. The second failure mode affected aerenchyma, "
            "where IoU spanned the full range from 0 to 0.97 (median 0.720). The "
            "lowest aerenchyma scores were concentrated in pearl millet, where "
            "individual lacunae are an order of magnitude smaller than in rice or "
            "sorghum and frequently fall below 50 pixels in area, near the lower "
            "bound at which any segmentation method can resolve them reliably. "
            "Both failure modes are therefore attributable to limits of the "
            "underlying image data, namely staining and sectioning artifacts in "
            "the case of outer-ring failures and pixel-scale resolution in the "
            "case of small aerenchyma lacunae, rather than to a structural deficit "
            "of the model architecture.",

            # Para 6: annotator-2 study
            "This interpretation was reinforced by an inter-annotator comparison. "
            "To assess whether the residual segmentation error reflects a "
            "fundamental ambiguity in the data rather than a limitation of the "
            "model, we asked a second expert plant biologist to re-annotate a "
            "balanced subset of 16 samples spanning all four species groups, "
            "blinded to the consensus ground-truth annotations used during "
            "training. Across these 16 samples, the second annotator agreed with "
            "the consensus ground truth at a mean IoU of 0.856, while RootSeg's "
            "predictions agreed with the same ground truth at a mean IoU of 0.893 "
            "(Fig. 2X). RootSeg matched or exceeded the human annotator on every "
            "individual tissue class. The largest gaps were on aerenchyma (model "
            "0.518, human 0.455), epidermis (model 0.830, human 0.747), and "
            "endodermis (model 0.922, human 0.870), all classes whose boundaries "
            "depend on subjective interpretation of weakly stained or "
            "developmentally variable cell-wall features. Per-species, RootSeg "
            "outperformed the second annotator on Sorghum (model 0.902, human "
            "0.830), Tomato (model 0.946, human 0.915), and Pearl Millet (model "
            "0.836, human 0.814), and was within 0.015 mIoU of the human annotator "
            "on Rice (model 0.823, human 0.837), where the small subsample of "
            "three samples made the comparison underpowered. Critically, "
            "aerenchyma was the lowest-scoring class for both the model and the "
            "second human annotator (IoU 0.518 and 0.455 respectively), confirming "
            "that the residual aerenchyma error reflects an irreducible labeling "
            "ambiguity rather than a model-specific limitation. RootSeg's "
            "predictions are therefore at least as reliable as a second expert's "
            "manual annotations on the same images, and the model has reached the "
            "upper bound of what can be achieved on this task without finer-"
            "grained annotation guidelines.",

            # Para 7: comparison setup (from user)
            "To identify which architectural choices drive RootSeg's performance, "
            "we trained six alternative models on the same dataset using the same "
            "loss combination and the same augmentation pipeline. Two were "
            "convolutional baselines built on ResNet34 and ResNet50 encoders "
            "pretrained on ImageNet and paired with a UNet++ decoder, "
            "representing the standard convolutional segmentation approach used "
            "in most prior plant-imaging work. Three were foundation-model "
            "variants that probed which part of our architecture mattered most. "
            "One paired the same DPT decoder with DINOv2, an older self-"
            "supervised foundation model trained on a smaller image corpus. The "
            "other two replaced the DPT decoder with simpler segmentation heads, "
            "namely a multi-scale linear head paired with DINOv2 and a small "
            "multi-layer perceptron paired with DINOv3. The sixth, YOLO26m-seg, "
            "was a state-of-the-art instance-segmentation model trained from a "
            "COCO-pretrained checkpoint, representing the per-object detection "
            "paradigm common in cell-segmentation packages.",

            # Para 8: comparison numbers (from user) + Fig 2c reference
            "On the in-distribution test set, all seven models scored within a "
            "narrow range of mIoU 0.875 to 0.889, indicating that the dataset is "
            "rich enough that several modern architectures can fit the training "
            "distribution well (Fig. 2f). The picture changed dramatically on the "
            "zero-shot Zeiss confocal split, where the spread of mIoU widened "
            "from 0.805 to 0.888. RootSeg (DINOv3 plus DPT) and the closely "
            "related DINOv2 plus DPT variant both retained mIoU 0.888, with no "
            "measurable degradation relative to the in-distribution split. The "
            "convolutional ResNet baselines dropped to 0.851 for ResNet50 and "
            "0.871 for ResNet34. The simpler-decoder foundation-model variants "
            "dropped to 0.863 (DINOv2 plus multi-scale linear head) and 0.856 "
            "(DINOv3 plus multi-layer perceptron head). The YOLO instance-"
            "segmentation baseline dropped sharpest, to 0.805. Visual examples "
            "of RootSeg's zero-shot Zeiss segmentations are shown in Fig. 2c, "
            "with mIoU values from 0.863 to 0.920, illustrating that the model "
            "retains sharp tissue boundaries on the unseen confocal optics.",

            # Para 9: comparison interpretation (from user)
            "These comparisons indicate that two architectural choices together "
            "enable RootSeg to generalize across imaging platforms. The first is "
            "self-supervised foundation-model pretraining, which exposes the "
            "encoder to a far broader range of visual statistics than ImageNet "
            "alone. The second is multi-scale dense prediction decoding, which "
            "preserves spatial detail at every resolution rather than only at "
            "the output. Convolutional encoders pretrained on a smaller, less "
            "diverse image collection appear to learn platform-specific visual "
            "cues that fail to transfer, and instance-segmentation models, which "
            "classify each lacuna independently and lack a coherent global "
            "tissue-map representation, transfer worst of all.",

            # Para 10: architecture rationale
            "We chose this hybrid architecture because root tissue identity is "
            "determined as much by global radial position as by local cellular "
            "texture. The exodermis, for example, is operationally defined as "
            "the second cell layer from the outer surface, and a model that "
            "classifies regions from local appearance alone cannot recover the "
            "correct label in samples where the exodermis is unlignified and "
            "visually indistinguishable from the surrounding cortex. The vision "
            "transformer encoder, which can attend to features anywhere in the "
            "image at every layer of processing, provides exactly the long-range "
            "context required to resolve such ambiguities. Self-supervised "
            "pretraining on a large natural-image collection further allows the "
            "encoder to learn broadly useful visual features such as texture, "
            "edge structure, and spatial coherence without those images needing "
            "to be labelled, and fine-tuning on our annotated root dataset "
            "adapts those features to the appearance of fluorescence "
            "microscopy.",

            # Para 11: brief training (full details in Methods)
            "We trained the encoder and decoder jointly on all 1,293 training "
            "samples. The pretrained encoder was fine-tuned at a low learning "
            "rate to preserve its pretrained features, while the randomly "
            "initialized decoder was trained at a higher learning rate to allow "
            "rapid learning of the task-specific cell-type-map output. The "
            "training objective combined four complementary loss terms designed "
            "to address the extreme class imbalance characteristic of root "
            "anatomy, where small structures such as aerenchyma can occupy as "
            "little as 0.5 percent of pixels in pearl millet. Geometric and "
            "intensity augmentations increased the effective dataset size, and "
            "channel-wise dropout promoted robustness to single-channel staining "
            "failures. Full hyperparameters and the exact loss specification are "
            "reported in Methods.",
        ],
    ),

    # ── Result 4 ───────────────────────────────────────────────────────────
    (
        "A unified model for monocots and dicots generalizes better than the "
        "specialist models trained on monocots or dicots individually",
        [
            # Para 1: motivation + experimental setup
            "Cereal cross-section images (rice, sorghum, pearl millet) and tomato "
            "cross-section images differ substantially in cell-wall architecture, "
            "root diameter, and the presence of aerenchyma. A natural concern when "
            "training a single model jointly on both clades is that the model "
            "might dilute its representation of either clade in service of the "
            "other, and that a clade-specific specialist would perform better on "
            "its own clade than a generalist. To test this directly, we trained "
            "two additional models. The Monocot Specialist was trained only on "
            "the 856 monocot training samples (rice, sorghum, and pearl millet). "
            "The Dicot Specialist was trained only on the 437 tomato training "
            "samples. The Unified Model and the two specialists shared an "
            "identical encoder, decoder, loss combination, augmentation pipeline, "
            "optimizer, and schedule, with the only difference being the species "
            "composition of the training set.",

            # Para 2: Fig 3a + 3b within-clade performance
            "We evaluated all three models on three test cohorts. The cohorts "
            "comprised the 148 monocot in-distribution test samples, the 37 dicot "
            "in-distribution test samples, and the 35 rice samples held out on "
            "the Zeiss confocal microscope. Representative prediction examples "
            "from each model on a monocot test sample (Millet, AW23 genotype), a "
            "dicot test sample (Tomato, slmyb92 mutant), and a Zeiss zero-shot "
            "sample (Rice, Kittake) are shown in Fig. 3a alongside the input "
            "image and the expert ground truth. The quantitative comparison "
            "across all 220 evaluation samples is summarized in Fig. 3b. On its "
            "own training clade, each specialist matched the Unified Model. On "
            "the 148 monocot test samples, the Monocot Specialist reached "
            "6-class mIoU 0.856 and the Unified Model reached 0.857, a "
            "difference of 0.001. On the 37 dicot test samples, the Dicot "
            "Specialist reached 6-class mIoU 0.932 and the Unified Model reached "
            "0.936, a difference of 0.004. On the 35 Zeiss zero-shot rice "
            "samples, both monocot-trained models reached 6-class mIoU 0.873 "
            "(Unified) and 0.872 (Monocot Specialist), again a difference of "
            "0.001. Joint training across clades therefore preserves full "
            "specialist-level accuracy on each clade rather than diluting it.",

            # Para 3: Off-clade collapse of specialists
            "The picture changed dramatically when each specialist was evaluated "
            "on the clade it had never seen during training (Fig. 3b, off-domain "
            "bars). The Monocot Specialist, applied to the 37 dicot test "
            "samples, dropped from 6-class mIoU 0.856 on its own monocot test "
            "set to 0.603 on tomato, a fall of 0.253. The Dicot Specialist, "
            "applied to the 148 monocot test samples, dropped from 0.932 on its "
            "own dicot test set to 0.589 on cereals, a fall of 0.343. On the "
            "Zeiss zero-shot split, which consists entirely of rice samples and "
            "is therefore off-domain for the Dicot Specialist, the Dicot "
            "Specialist collapsed further to 6-class mIoU 0.401, an absolute "
            "drop of 0.531 from its dicot in-distribution accuracy. By "
            "contrast, the Unified Model maintained 6-class mIoU between 0.857 "
            "and 0.936 across every cohort. The visual prediction panels in "
            "Fig. 3a illustrate this failure mode directly. The off-clade "
            "specialist predictions either miss entire tissue domains or "
            "misclassify large regions as the wrong tissue, while the Unified "
            "Model's prediction is essentially indistinguishable from the "
            "expert ground truth on both clades. Specialist training therefore "
            "produces a brittle model that fails catastrophically when applied "
            "beyond its training clade, whereas joint training produces a "
            "generalist model with no measurable performance cost.",

            # Para 4: PCA setup — framed as fine-tuning sanity check
            "We also asked whether fine-tuning on our annotated dataset "
            "reshapes the encoder's internal representation in a way that "
            "reflects root anatomy, rather than only modifying the downstream "
            "segmentation head. To answer this, we visualized the spatial "
            "structure of the encoder's learned representations directly on "
            "the input image (Fig. 3c). For each sample, we passed the image "
            "through the encoder and extracted the grid of patch-token "
            "feature vectors output by the final transformer block, then "
            "projected these high-dimensional vectors into a three-dimensional "
            "colour space using principal component analysis (PCA). PCA is a "
            "standard linear dimensionality-reduction method that identifies "
            "the directions of greatest variance in a high-dimensional "
            "feature set and projects the data onto those directions. When "
            "applied to patch-token features, PCA recovers the dominant "
            "patterns of variation across patches within an image. We "
            "rendered the first three principal components as the red, green, "
            "and blue channels of an output image at the patch grid "
            "resolution, with PCA computed separately for the pretrained "
            "encoder and for the fine-tuned Unified Model encoder (full "
            "procedure in Methods).",

            # Para 5: Fig 3c interpretation
            "Fig. 3c shows the encoder PCA visualization for three "
            "representative samples, namely a monocot test image (Rice, "
            "wox10-15 mutant), a dicot test image (Tomato, slmyb92 mutant), "
            "and a Zeiss zero-shot monocot image (Rice, Kittake), each "
            "rendered under two encoder states. The first column shows the "
            "pretrained DINOv3 encoder before any fine-tuning on our dataset, "
            "and the second column shows the fine-tuned Unified Model encoder. "
            "In the pretrained encoder, the PCA colours form diffuse and "
            "largely amorphous patterns within the root that do not align with "
            "the underlying anatomical tissue domains. In the Unified Model "
            "encoder, by contrast, the PCA colours form sharp, spatially "
            "coherent rings and regions that align directly with the "
            "epidermis, cortex, endodermis, exodermis, and vascular cylinder "
            "visible in the input image. Concentric ring structure emerges in "
            "the encoder feature space that mirrors the radial anatomical "
            "structure of the root, and aerenchyma regions are picked out as "
            "distinct colour patches embedded within the cortex. This pattern "
            "holds across all three rows, including the Zeiss zero-shot row in "
            "which the encoder had never seen confocal imaging optics during "
            "training.",

            # Para 6: synthesis — fine-tuning worked, encoder learned anatomy
            "This comparison confirms that fine-tuning on our annotated root "
            "dataset successfully reshapes the encoder's representation. "
            "Before fine-tuning, the pretrained DINOv3 encoder produces a "
            "feature space whose dominant axes of variation do not correspond "
            "to root anatomical structure. After fine-tuning, the encoder "
            "learns a more structured feature space that is correlated with "
            "the radial anatomical organization of the root, with each tissue "
            "domain occupying a distinct region of the feature space. The "
            "encoder is therefore not merely supporting a downstream "
            "segmentation head, it is learning representations that are "
            "themselves anatomically meaningful.",
        ],
    ),

    # ── Methods ────────────────────────────────────────────────────────────
    (
        "Methods",
        [
            {"sub": "Plant materials and growth conditions"},
            {"placeholder":
             "[TBD by co-author. Should describe genotype list per species, "
             "growth substrate (e.g. agar plates, soil, hydroponics), growth "
             "conditions (temperature, photoperiod), and harvest stage and "
             "developmental window for each species and experiment.]"},

            {"sub": "Tissue preparation and fluorescent staining"},
            {"placeholder":
             "[TBD by co-author. Should describe root excision, fixation, "
             "sectioning method (vibratome thickness, blade), tissue clearing "
             "protocol (e.g. ClearSee), and the fluorescent staining recipe "
             "for each dye, including concentration and incubation conditions.]"},

            {"sub": "Image acquisition"},
            "Cross-sections were imaged on three platforms. The Olympus IX83 "
            "widefield microscope was used for the majority of samples, with "
            "[objective, exposure, channel filter set TBD by co-author]. The "
            "Cytation C10 plate-reader imager (BioTek) was used for a subset of "
            "rice, sorghum, and tomato samples, with [imaging mode and channel "
            "settings TBD by co-author]. The Zeiss LSM 970 confocal microscope "
            "was used to acquire the held-out zero-shot test cohort of 35 rice "
            "samples, with [laser lines, excitation and emission ranges, "
            "pinhole, and z-stack handling TBD by co-author]. Across all "
            "platforms, three channels were collected per sample (DAPI, FITC, "
            "TRITC) and saved as separate single-channel grayscale TIFF files. "
            "Original pixel resolution and field of view varied across "
            "platforms, and platform-specific characteristics were preserved "
            "through to the model input (see Dataset construction and splits).",

            {"sub": "Annotation protocol"},
            "Each cross-section was annotated by a trained plant biologist "
            "using a custom polygon-editing tool implemented in PyQt5. The "
            "tool overlaid the three fluorescence channels with adjustable "
            "brightness and gamma controls and supported brush-mode region "
            "drawing, polygon vertex editing, and ring splitting. Annotators "
            "drew six classes of polygons. These were the whole root contour "
            "(one polygon per sample), aerenchyma regions (zero or more "
            "polygons per sample), the outer and inner contours of the "
            "endodermis (one each per sample), and the outer and inner "
            "contours of the exodermis (one each per sample, biologically "
            "meaningful only in tomato). Ring-tissue boundaries were placed "
            "along the outer edge of their respective cell walls, with the "
            "outer-minus-inner construction used to recover ring thickness "
            "downstream. All polygons were saved in YOLO normalized polygon "
            "format with a class identifier and (x, y) vertex coordinates. The "
            "total annotation effort across the dataset was 31,595 polygons "
            "across 1,695 samples. [TBD by co-author. Annotator headcount and "
            "approximate time investment.]",

            {"sub": "Dataset construction and splits"},
            "The 1,695 samples were partitioned into a training set (1,293 "
            "samples), a validation set (182 samples), an in-distribution test "
            "set (185 samples), and a zero-shot Zeiss test set (35 samples). "
            "Splits were assigned at the experiment level rather than at the "
            "sample level, ensuring that no two samples from a single biological "
            "experiment appeared in different splits. The training, validation, "
            "and in-distribution test sets together comprised samples from rice, "
            "sorghum, pearl millet, and tomato imaged on the Olympus IX83 and "
            "Cytation C10 platforms. The zero-shot test set comprised 35 rice "
            "cross-sections imaged on the Zeiss LSM 970 confocal microscope, a "
            "platform not represented in the training data. Each input image "
            "was constructed by stacking the three single-channel TIFF files "
            "(DAPI, FITC, TRITC) along the channel dimension and resampling to "
            "1,024 \u00d7 1,024 pixels with bilinear interpolation. Channel "
            "intensities were normalized per image by clipping to the 1st and "
            "99.5th percentile of pixel values within that channel and "
            "rescaling to the unit interval, producing a (3, 1024, 1024) "
            "float32 tensor per sample.",

            {"sub": "Model architecture (RootSeg)"},
            "RootSeg comprised a self-supervised vision transformer encoder "
            "(DINOv3-S/16, Hugging Face checkpoint "
            "facebook/dinov3-vits16-pretrain-lvd1689m) and a custom Meta-style "
            "dense prediction transformer (DPT) decoder. The encoder accepted "
            "a (3, 1024, 1024) input tensor and partitioned the image into "
            "16 \u00d7 16 non-overlapping patches, producing a token sequence "
            "of length 1 plus 4096 (one [CLS] token plus 64 \u00d7 64 spatial "
            "tokens). Each token had embedding dimension 384. The encoder "
            "applied 12 transformer blocks of self-attention and feed-forward "
            "MLP layers to the token sequence. The DPT decoder tapped the "
            "patch-token output of layers 3, 6, 9, and 12 of the encoder "
            "(four evenly spaced layers spanning shallow to deep features), "
            "discarded the [CLS] token, and reshaped the remaining tokens to "
            "a (384, 64, 64) feature map per tap. Each tap was projected to a "
            "multi-resolution pyramid using transposed-convolution upsampling "
            "and 1 \u00d7 1 convolutional projections, producing feature maps "
            "at strides 4, 8, 16, and 32 with channel widths 96, 192, 384, "
            "and 768, respectively. These four feature maps were progressively "
            "fused from deepest to shallowest through residual convolutional "
            "blocks at a unified width of 256 channels, with bilinear "
            "upsampling at each fusion step. The residual-fusion design "
            "serves two complementary purposes. Skip connections from the "
            "encoder to the decoder at every resolution stabilize learning, "
            "because each decoder block only needs to learn an incremental "
            "refinement of the existing feature representation rather than "
            "reconstruct it from scratch, and the multi-scale pyramid ensures "
            "that the final per-pixel tissue prediction is informed both by "
            "the precise local cellular boundaries visible in the shallow "
            "features and by the long-range anatomical context encoded in the "
            "deep features. The final fused feature map at stride 4 was "
            "passed through a final upsampling stage and a 1 \u00d7 1 "
            "convolutional segmentation head to produce a (7, 1024, 1024) "
            "class-probability tensor over the seven tissue classes. Argmax "
            "across the class dimension produced the discrete tissue map.",

            {"sub": "Training procedure"},
            "RootSeg was trained on a single NVIDIA H100 80 GB GPU using "
            "PyTorch 2 and PyTorch Lightning. Training used bfloat16 "
            "mixed-precision throughout, except for the Lov\u00e1sz-softmax "
            "loss term, which was computed in float32 within an explicit "
            "autocast(enabled=False) block to avoid a precision mismatch "
            "between the float16 logits and the float32 sub-gradient "
            "computation in segmentation_models_pytorch. The optimizer was "
            "AdamW with weight decay 0.01 and a differential learning rate. "
            "The encoder was fine-tuned at learning rate 1 \u00d7 10\u207b\u2075 "
            "and the randomly initialized decoder was trained at learning "
            "rate 1 \u00d7 10\u207b\u2074, both decayed via a cosine "
            "annealing schedule across the full training budget. Training "
            "used a per-GPU batch size of 16 samples, a maximum of 200 "
            "epochs, and early stopping on validation loss with patience 15 "
            "epochs. The composite training objective was a sum of four loss "
            "terms applied to the full 7-class softmax output. The Dice loss "
            "(smp.losses.DiceLoss, multiclass mode) maximizes per-class "
            "region overlap and is robust to extreme class imbalance because "
            "it weights all classes by their predicted-versus-true region "
            "overlap rather than by pixel count. The focal loss "
            "(smp.losses.FocalLoss, multiclass mode, default \u03b3 = 2) "
            "downweights well-classified pixels, redirecting gradient toward "
            "the hard and rare regions that dominate the residual error, "
            "such as small aerenchyma lacunae. The cross-entropy loss "
            "(F.cross_entropy with equal class weights of 1.0 across all "
            "seven classes) is the standard per-pixel objective for semantic "
            "segmentation. The Lov\u00e1sz-softmax loss "
            "(smp.losses.LovaszLoss, multiclass mode) is a differentiable "
            "surrogate for the IoU metric that produces sharper class "
            "boundaries than cross-entropy. The four loss terms contributed "
            "equally to the total. The random seed was fixed at 42.",

            {"sub": "Augmentation pipeline"},
            "Augmentation was applied online during training using the "
            "albumentations library, with a single pipeline shared across all "
            "segmentation models compared in this study to keep architectural "
            "comparisons fair. The geometric augmentations were RandomRotate90 "
            "(p = 0.5), horizontal flip (p = 0.5), vertical flip (p = 0.5), "
            "an Affine transform combining translation, scaling, rotation, "
            "and shear (p = 0.5), and ElasticTransform (p = 0.3). The "
            "intensity augmentations were RandomBrightnessContrast (p = 0.5), "
            "RandomGamma (p = 0.5), GaussianBlur (p = 0.3), and GaussNoise "
            "(p = 0.3). To promote robustness to single-channel staining "
            "failures and to inter-platform channel intensity differences, "
            "ChannelDropout (p = 0.2, fill 0) and ChannelShuffle (p = 0.2) "
            "were applied to the three input channels. Hue and saturation "
            "jittering were not applied because the three fluorescence "
            "channels carry fixed biological meaning rather than perceptual "
            "color. After all augmentations, the image was resized to "
            "1,024 \u00d7 1,024.",

            {"sub": "Comparison architectures"},
            "The six comparison architectures were trained with identical "
            "data, loss combination, optimizer, schedule, batch size, and "
            "epoch budget, with the only differences being the encoder, the "
            "decoder, and (where required by the architecture) the "
            "segmentation head. The convolutional baselines were two "
            "configurations of the segmentation_models_pytorch UNet++ decoder "
            "paired with ImageNet-pretrained ResNet34 (24.4 M parameters) and "
            "ResNet50 (32.5 M parameters) encoders. The two simpler-decoder "
            "variants paired DINOv2 (vit_small_patch14_dinov2 from timm, "
            "pretrained on LVD-142M) with a multi-scale linear segmentation "
            "head (Meta MS-Linear recipe, taps the last four transformer "
            "blocks) and DINOv3-S/16 with a small per-patch MLP segmentation "
            "head. The DINOv2-DPT variant used the same DPT-meta decoder as "
            "RootSeg paired with the DINOv2 encoder. The instance-"
            "segmentation baseline was YOLO26m (Ultralytics) trained from a "
            "COCO-pretrained checkpoint on the 6 raw classes (whole root, "
            "aerenchyma, outer and inner endodermis, outer and inner "
            "exodermis), with overlap_mask=False to ensure that nested "
            "polygons were not flattened into ring-only masks at training "
            "time. After inference, instance polygon predictions were merged "
            "into pixel-level semantic masks using the same outer-minus-inner "
            "ring construction described under Evaluation, allowing direct "
            "comparison with the semantic models on the same Bio-7 metric.",

            {"sub": "Evaluation metrics"},
            "Performance was measured on the held-out in-distribution test "
            "set (185 samples) and the zero-shot Zeiss test set (35 samples). "
            "The primary metrics were per-class intersection-over-union (IoU) "
            "and Dice coefficient. For a predicted binary mask P and a "
            "ground-truth binary mask G of the same tissue class, the IoU is "
            "defined as the number of pixels where both P and G are positive "
            "(the intersection) divided by the number of pixels where either "
            "P or G is positive (the union). The Dice coefficient is defined "
            "as twice the intersection divided by the sum of the two mask "
            "areas, equivalent to 2 times IoU divided by (1 plus IoU). Both "
            "metrics range from 0 (no overlap) to 1 (perfect overlap). Both "
            "metrics were computed on the raw model output without any "
            "post-processing or polygon round-tripping. For all models, the "
            "predicted output was mapped to seven biologically interpretable "
            "tissue domains (whole root, epidermis, exodermis, cortex, "
            "endodermis, vascular cylinder, aerenchyma) using a deterministic "
            "pixel-level rule. For the semantic models, the model directly "
            "predicts these seven classes. For the YOLO instance-segmentation "
            "baseline, ring tissues were derived from the raw 6-class polygon "
            "predictions by computing pixel-level differences between outer "
            "and inner ring masks (the endodermis ring as outer-endodermis "
            "minus inner-endodermis, the exodermis ring analogously). The "
            "vascular cylinder was defined as the interior of the inner-"
            "endodermis polygon, and the cortex was defined as the whole-root "
            "mask minus the union of all other tissue masks. The same "
            "conversion was applied identically to ground-truth labels and "
            "predictions, so any conversion artifact would cancel out when "
            "comparing the two. Per-class IoU was averaged over samples in "
            "which the class was present. For aerenchyma in tomato, where "
            "the class is biologically absent, IoU was undefined and excluded "
            "from per-class averages. The mean IoU and mean Dice reported "
            "per sample were averaged across the seven Bio-7 classes for "
            "that sample.",

            {"sub": "Specialist models (Monocot Specialist and Dicot Specialist)"},
            "Two clade-specialist models were trained for the unified-versus-"
            "specialist comparison in Result 4. The Monocot Specialist was "
            "trained on the 856 training samples drawn from rice, sorghum, and "
            "pearl millet. The Dicot Specialist was trained on the 437 training "
            "samples drawn from tomato (including the cultivated Solanum "
            "lycopersicum and 13 wild Solanum species). The validation set was "
            "restricted to the matching clade for each specialist. Both "
            "specialists used the same DINOv3-S/16 encoder, Meta-style DPT "
            "decoder, four-component composite loss, augmentation pipeline, "
            "optimizer, learning-rate schedule, batch size, and epoch budget "
            "as the Unified Model, with the only difference being the "
            "training-set species composition. Specialists were evaluated on "
            "three cohorts. Each was evaluated on its own clade's "
            "in-distribution test split using the standard "
            "metrics_bio7.csv pipeline. Each was also evaluated on the "
            "off-clade test split (the Monocot Specialist on the 37 tomato "
            "test samples, the Dicot Specialist on the 148 cereal test "
            "samples) using a separate off-domain inference pass written to a "
            "metrics_bio7_offdomain.csv file. Finally, both specialists were "
            "applied zero-shot to the 35 Zeiss rice samples for direct "
            "comparison with the Unified Model on the zero-shot split.",

            {"sub": "Inter-annotator comparison"},
            "To estimate the upper bound on segmentation accuracy that any "
            "model trained on this dataset could be expected to achieve, we "
            "asked a second expert plant biologist to re-annotate a subset "
            "of 16 samples drawn from the in-distribution test set. The "
            "subset was balanced across the four species groups (six "
            "Sorghum, five Tomato, three Rice, two Pearl Millet) and across "
            "the two widefield platforms (14 Olympus, 2 Cytation C10). The "
            "second annotator received the raw three-channel images only and "
            "was blinded to the original consensus annotations used as model "
            "ground truth. Per-class and per-sample IoU were computed "
            "between the second annotator's polygons and the consensus "
            "ground truth using the same Bio-7 conversion and the same IoU "
            "definition described above. Per-class IoU was undefined and "
            "excluded for any class that was absent in both the consensus "
            "and the second annotator (this affected only aerenchyma in the "
            "five Tomato samples).",

            {"sub": "Encoder feature visualization by principal component analysis"},
            "To visualize the spatial organization of the encoder's learned "
            "representations on a per-sample basis, we used a two-pass "
            "principal component analysis (PCA) procedure adapted from the "
            "DINOv2 and DINOv3 feature-visualization conventions (Oquab et al. "
            "2023, Sim\u00e9oni et al. 2025). For each sample, we resampled "
            "the input image to 1,536 \u00d7 1,536 pixels (slightly above the "
            "standard 1,024 \u00d7 1,024 training resolution to give the "
            "visualization a finer patch grid), passed it through the chosen "
            "encoder (either the frozen pretrained DINOv3-S/16 or the "
            "fine-tuned Unified Model encoder), and extracted the patch-token "
            "output of the final transformer block. With a patch size of 16, "
            "this produced a 96 \u00d7 96 grid of 384-dimensional patch-token "
            "feature vectors per sample.",

            "The first PCA pass fit a one-component PCA on the patch tokens of "
            "a single image. The first principal component of self-supervised "
            "DINO-family features is known to separate foreground (object) "
            "from background (surroundings) without supervision. We "
            "thresholded the PC1 scores at their median to obtain a binary "
            "foreground mask, choosing the foreground side as the one with "
            "fewer image-corner patches under the assumption that the four "
            "corners of a cross-section image are background.",

            "The second PCA pass fit a three-component PCA on the union of "
            "foreground tokens pooled across all samples for the encoder "
            "being visualized, with each sample's foreground tokens "
            "mean-centered before pooling. The pretrained DINOv3 encoder and "
            "the fine-tuned Unified Model encoder used separate PCA bases, "
            "each fit only on its own foreground tokens, because the two "
            "encoders live in different feature spaces and a shared PCA "
            "basis would not be meaningful between them. For each (encoder, "
            "sample) pair, the three component scores at each foreground "
            "patch were clipped at their per-component 2nd and 98th "
            "percentile values (pooled across the fit), then linearly "
            "rescaled to the unit interval and assigned to the red (PC1), "
            "green (PC2), and blue (PC3) channels of an output image at the "
            "patch grid resolution. Background patches were rendered black. "
            "The result is a per-sample colour visualization in which "
            "patches with similar encoder representations appear in similar "
            "colours, so the spatial layout of those colours directly "
            "reflects the structure encoded in the encoder's feature space.",

            {"sub": "Downstream measurements"},
            "Seven biology-relevant per-sample measurements were computed "
            "from each predicted tissue map and from the corresponding "
            "ground-truth tissue map. The aerenchyma area ratio was computed "
            "as the total pixel count of the aerenchyma class divided by the "
            "pixel count of the whole-root mask. Mean intensity measurements "
            "were computed by masking the un-normalized TRITC channel "
            "image with the predicted (or ground-truth) tissue mask and "
            "taking the mean pixel intensity within the mask, then repeating "
            "the procedure with the FITC channel. This was done independently "
            "for the exodermis, endodermis, and vascular cylinder classes, "
            "yielding two intensity readouts per tissue. Predicted versus "
            "ground-truth measurements were compared by ordinary "
            "least-squares linear regression, reporting the coefficient of "
            "determination (R\u00b2) and the slope per measurement, "
            "computed both pooled across species and stratified by species.",

            {"sub": "Code and data availability"},
            {"placeholder":
             "[TBD. Source code for training, evaluation, downstream analysis, "
             "and the polygon-editing tool is available at GitHub URL. The "
             "annotated dataset and pretrained model weights are deposited at "
             "data repository under DOI.]"},
        ],
    ),
]

# ── Build doc ────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Helvetica"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Results 1, Results 2, Results 4, and Methods (revised draft, 2026-05-10)")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

audit_text_chunks = []

for heading, items in SECTIONS:
    h = doc.add_paragraph()
    h_run = h.add_run(heading)
    h_run.bold = True
    h_run.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)

    for item in items:
        if isinstance(item, dict) and "sub" in item:
            sh = doc.add_paragraph()
            sh_run = sh.add_run(item["sub"])
            sh_run.bold = True
            sh_run.font.size = Pt(11)
            sh.paragraph_format.space_before = Pt(8)
            sh.paragraph_format.space_after = Pt(2)
            audit_text_chunks.append(item["sub"])
        elif isinstance(item, dict) and "placeholder" in item:
            pp = doc.add_paragraph()
            pp_run = pp.add_run(item["placeholder"])
            pp_run.italic = True
            pp_run.font.color.rgb = PLACEHOLDER
            pp.paragraph_format.first_line_indent = Inches(0.25)
            pp.paragraph_format.space_after = Pt(8)
        else:
            body = doc.add_paragraph(item)
            body.paragraph_format.first_line_indent = Inches(0.25)
            body.paragraph_format.space_after = Pt(8)
            audit_text_chunks.append(item)

audit_text = "\n".join(audit_text_chunks)
banned = {"\u2014": "em dash", ";": "semicolon", ":": "colon"}
for ch, name in banned.items():
    if ch in audit_text:
        idx = audit_text.index(ch)
        ctx = audit_text[max(0, idx - 50) : idx + 50]
        raise AssertionError(f"{name} found in body prose near: ...{ctx}...")

doc.save(OUT)
n_words = sum(
    len(item.split()) for _, items in SECTIONS for item in items
    if isinstance(item, str)
)
n_placeholder_words = sum(
    len(item["placeholder"].split())
    for _, items in SECTIONS for item in items
    if isinstance(item, dict) and "placeholder" in item
)
print(f"Wrote {OUT}")
print(f"Body word count (drafted prose): {n_words}")
print(f"Placeholder word count (TBD by co-author): {n_placeholder_words}")
